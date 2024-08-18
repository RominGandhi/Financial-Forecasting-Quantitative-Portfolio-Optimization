import os
import numpy as np
import pandas as pd
import json
from datetime import datetime
from data_fetcher import fetch_data, get_alpha_vantage_asset_name, get_asset_name, get_cryptocompare_asset_name, save_data
from portfolio_metrics import calculate_portfolio_metrics
from optimization import optimize_portfolio
from docx import Document
import subprocess
import platform

# Directory where user logs will be saved
USER_LOG_DIR = os.path.join(os.getcwd(), 'user_log')
os.makedirs(USER_LOG_DIR, exist_ok=True)

# Directory for config files
CONFIG_DIR = os.path.join(os.getcwd(), 'data', 'config')
os.makedirs(CONFIG_DIR, exist_ok=True)

def create_word_doc(results, doc_filename):
    tickers = results['tickers']
    start_date = results['start_date']
    end_date = results['end_date']
    initial_metrics = results['initial_metrics']
    optimal_weights = results['optimal_weights']
    optimal_metrics = results['optimal_metrics']

    # Create a new Document
    doc = Document()
    doc.add_heading('Portfolio Analysis and Optimization Results', 0)

    # Add configuration details
    doc.add_heading('Configuration', level=1)
    doc.add_paragraph(f"Tickers: {', '.join(tickers)}")
    doc.add_paragraph(f"Start Date: {start_date}")
    doc.add_paragraph(f"End Date: {end_date}")

    # Add initial portfolio metrics
    doc.add_heading('Initial Portfolio Metrics', level=1)
    doc.add_paragraph(f"Expected Return: {initial_metrics['return'] * 100:.2f}%")
    doc.add_paragraph(f"Volatility: {initial_metrics['volatility'] * 100:.2f}%")
    doc.add_paragraph(f"Sharpe Ratio: {initial_metrics['sharpe_ratio']:.2f}")

    # Add optimization results
    doc.add_heading('Optimized Portfolio Metrics', level=1)
    doc.add_paragraph('Optimal Weights:')
    for ticker, weight in zip(tickers, optimal_weights):
        doc.add_paragraph(f"{ticker}: {weight * 100:.2f}%")

    doc.add_paragraph(f"Optimized Expected Return: {optimal_metrics['return'] * 100:.2f}%")
    doc.add_paragraph(f"Optimized Volatility: {optimal_metrics['volatility'] * 100:.2f}%")
    doc.add_paragraph(f"Optimized Sharpe Ratio: {optimal_metrics['sharpe_ratio']:.2f}")

    # Save the document
    doc.save(doc_filename)

def open_word_doc(doc_filename):
    if platform.system() == 'Darwin':       # macOS
        subprocess.call(('open', doc_filename))
    elif platform.system() == 'Windows':    # Windows
        os.startfile(doc_filename)
    elif platform.system() == 'Linux':      # Linux
        subprocess.call(('xdg-open', doc_filename))

def save_tickers_config(tickers, start_date, end_date):
    """Save the tickers, start date, and end date to a JSON file in the config directory."""
    config_path = os.path.join(CONFIG_DIR, 'tickers.json')
    config_data = {
        "tickers": tickers,
        "start_date": start_date,
        "end_date": end_date
    }
    with open(config_path, 'w') as f:
        json.dump(config_data, f, indent=4)
    print(f"Configuration saved to {config_path}")

def main():
    # Ask user for input
    tickers_input = input("\nEnter the tickers separated by commas (e.g., AAPL,IEF.BOND,BTC.CRYPTO,SPY.ETF): ")
    tickers = [ticker.strip().upper() for ticker in tickers_input.split(",")]

    # Categorize and fetch full names for the assets
    print("\nAsset Names:")
    for ticker in tickers:
        try:
            if ticker.endswith('.BOND'):
                ticker_base = ticker.replace('.BOND', '')
                full_name = get_alpha_vantage_asset_name(ticker_base)
                if not full_name:
                    raise ValueError(f"Financial instrument {ticker} not found.")
                print(f"{ticker}: {full_name} (Bond)")
            elif ticker.endswith('.CRYPTO'):
                ticker_base = ticker.replace('.CRYPTO', '')
                full_name = get_cryptocompare_asset_name(ticker_base)
                if not full_name:
                    raise ValueError(f"Financial instrument {ticker} not found.")
                print(f"{ticker}: {full_name} (Cryptocurrency)")
            elif ticker.endswith('.ETF'):
                ticker_base = ticker.replace('.ETF', '')
                full_name = get_asset_name(ticker_base)
                if not full_name:
                    raise ValueError(f"Financial instrument {ticker} not found.")
                print(f"{ticker}: {full_name} (ETF)")
            else:
                full_name = get_asset_name(ticker)
                if not full_name:
                    raise ValueError(f"Financial instrument {ticker} not found.")
                print(f"{ticker}: {full_name} (Stock)")
        except Exception as e:
            print(f"Error: {e}. Exiting program.")
            return  # Exit the program

    start_date = input("\nEnter the start date (YYYY-MM-DD): ")
    end_date = input("Enter the end date (YYYY-MM-DD): ")

    # Save the tickers, start date, and end date to the config file
    save_tickers_config(tickers, start_date, end_date)

    # Fetch data
    print("Fetching data...")
    data = fetch_data(tickers, start_date, end_date)

    # Ask for a filename to save the data
    filename = input("Enter the filename to save the data (e.g., 'portfolio_data.csv'): ")
    save_data(data, filename)

    # Calculate daily returns
    returns = data.pct_change().dropna()

    # Ask the user for weights
    weights = []
    for ticker in tickers:
        while True:
            try:
                weight = float(input(f"Enter the weight for {ticker} (as a fraction, e.g., 0.2 for 20%): "))
                weights.append(weight)
                break
            except ValueError:
                print("Invalid input. Please enter a number.")

    # Ensure the weights sum to 1
    if not np.isclose(sum(weights), 1.0):
        print("Error: The sum of the weights must equal 1 (i.e., 100%). Please enter weights that sum to 1.")
        return  # Exit the program

    # Calculate portfolio metrics
    weights = np.array(weights)
    print("Calculating portfolio metrics...")
    metrics = calculate_portfolio_metrics(weights, returns)

    # Display portfolio metrics
    print(f"Portfolio Expected Return: {metrics['return'] * 100:.2f}%")
    print(f"Portfolio Volatility: {metrics['volatility'] * 100:.2f}%")
    print(f"Portfolio Sharpe Ratio: {metrics['sharpe_ratio']:.2f}")

    # Optimize portfolio
    print("Optimizing portfolio...")
    optimization_result = optimize_portfolio(returns, method='sharpe')
    optimal_weights = optimization_result['weights']
    optimal_metrics = optimization_result['metrics']

    # Display optimized portfolio results
    print("\nOptimal Weights:")
    for ticker, weight in zip(tickers, optimal_weights):
        print(f"{ticker}: {weight:.2%}")

    print(f"\nOptimized Portfolio Expected Return: {optimal_metrics['return']*100:.2f}%")
    print(f"Optimized Portfolio Volatility: {optimal_metrics['volatility']*100:.2f}%")
    print(f"Optimized Portfolio Sharpe Ratio: {optimal_metrics['sharpe_ratio']:.2f}")

    # Save results to a JSON file with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    json_filename = os.path.join(USER_LOG_DIR, f"portfolio_results_{timestamp}.json")
    results = {
        "tickers": tickers,
        "start_date": start_date,
        "end_date": end_date,
        "initial_metrics": metrics,
        "optimal_weights": optimal_weights.tolist(),
        "optimal_metrics": optimal_metrics
    }
    with open(json_filename, 'w') as f:
        json.dump(results, f)
    print(f"Results saved to {json_filename}")

    # Ask if the user wants to generate a Word document
    create_doc = input("Would you like to generate a Word document with the results? (yes/no): ").strip().lower()
    if create_doc == 'yes':
        doc_filename = os.path.join(USER_LOG_DIR, f"portfolio_results_{timestamp}.docx")
        create_word_doc(results, doc_filename)
        print(f"Word document saved to {doc_filename}")
        open_word_doc(doc_filename)

if __name__ == "__main__":
    main()
