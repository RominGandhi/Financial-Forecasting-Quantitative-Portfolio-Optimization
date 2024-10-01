import json
from docx import Document

def create_word_doc(results_filename, doc_filename):
    # Load results from JSON file
    with open(results_filename, 'r') as f:
        results = json.load(f)
    
    tickers = results['tickers']
    start_date = results['start_datae']
    end_date = results['end_date']
    initial_metrics = results['initial_metrics']
    optimal_weights = results['optimal_weights']
    optimal_metrics = results['optimal_metrics']

    # Create a new Document
    doc = Document()
    doc.add_heading('Portfolio Analysis and Optimization Results', 0)

    # Add configuration details
    doc.add_heading('Configuration', level=1)
    doc.add_paragraph(f"Tickers: {', '.join(tickers)}")
    doc.add_paragraph(f"Start Date: {start_date}")
    doc.add_paragraph(f"End Date: {end_date}")

    # Add initial portfolio metrics
    doc.add_heading('Initial Portfolio Metrics', level=1)
    doc.add_paragraph(f"Expected Return: {initial_metrics['return'] * 100:.2f}%")
    doc.add_paragraph(f"Volatility: {initial_metrics['volatility'] * 100:.2f}%")
    doc.add_paragraph(f"Sharpe Ratio: {initial_metrics['sharpe_ratio']:.2f}")

    # Add optimization results
    doc.add_heading('Optimized Portfolio Metrics', level=1)
    doc.add_paragraph('Optimal Weights:')
    for ticker, weight in zip(tickers, optimal_weights):
        doc.add_paragraph(f"{ticker}: {weight * 100:.2f}%")

    doc.add_paragraph(f"Optimized Expected Return: {optimal_metrics['return'] * 100:.2f}%")
    doc.add_paragraph(f"Optimized Volatility: {optimal_metrics['volatility'] * 100:.2f}%")
    doc.add_paragraph(f"Optimized Sharpe Ratio: {optimal_metrics['sharpe_ratio']:.2f}")

    # Save the document
    doc.save(doc_filename)

if __name__ == "__main__":
    results_filename = "portfolio_results.json"
    doc_filename = input("Enter the filename to save the Word document (e.g., 'portfolio_results.docx'): ")
    create_word_doc(results_filename, doc_filename)
