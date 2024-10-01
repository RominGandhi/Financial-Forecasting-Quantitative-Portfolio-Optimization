import numpy as np
import pandas as pd
import json
import os
from datetime import datetime
from data_fetcher import fetch_data, save_data, get_asset_name
from portfolio_metrics import calculate_portfolio_metrics
from optimization import optimize_portfolio
from docx import Document
import subprocess
import platform
import matplotlib.pyplot as plt
from predict_prices import load_historical_data, build_lstm_model, train_lstm_model, predict_future_prices_lstm, get_exchange_for_ticker

# Define the directory where user logs will be saved
USER_LOG_DIR = os.path.join(os.getcwd(), 'user_log')
os.makedirs(USER_LOG_DIR, exist_ok=True)

# Setting the default path for the historical data directory
DATA_DIR = os.path.join(os.getcwd(), 'data', 'historical_data')

# Path to the directory where predicted prices CSV files will be saved
PREDICTED_DIR = os.path.join(os.getcwd(), 'data', 'Predicted Prices')


def create_word_doc(results, doc_filename):
    tickers = results['tickers']
    start_date = results['start_date']
    end_date = results['end_date']
    initial_metrics = results['initial_metrics']
    optimal_weights = results['optimal_weights']
    optimal_metrics = results['optimal_metrics']

    # Create a new Document
    doc = Document()
    doc.add_heading('Portfolio Analysis and Optimization Results', 0)

    # Add configuration details
    doc.add_heading('Configuration', level=1)
    doc.add_paragraph(f"Tickers: {', '.join(tickers)}")
    doc.add_paragraph(f"Start Date: {start_date}")
    doc.add_paragraph(f"End Date: {end_date}")

    # Add initial portfolio metrics
    doc.add_heading('Initial Portfolio Metrics', level=1)
    doc.add_paragraph(f"Expected Return: {initial_metrics['return'] * 100:.2f}%")
    doc.add_paragraph(f"Volatility: {initial_metrics['volatility'] * 100:.2f}%")
    doc.add_paragraph(f"Sharpe Ratio: {initial_metrics['sharpe_ratio']:.2f}")

    # Add optimization results
    doc.add_heading('Optimized Portfolio Metrics', level=1)
    doc.add_paragraph('Optimal Weights:')
    for ticker, weight in zip(tickers, optimal_weights):
        doc.add_paragraph(f"{ticker}: {weight * 100:.2f}%")

    doc.add_paragraph(f"Optimized Expected Return: {optimal_metrics['return'] * 100:.2f}%")
    doc.add_paragraph(f"Optimized Volatility: {optimal_metrics['volatility'] * 100:.2f}%")
    doc.add_paragraph(f"Optimized Sharpe Ratio: {optimal_metrics['sharpe_ratio']:.2f}")

    # Save the document
    doc.save(doc_filename)


def open_word_doc(doc_filename):
    if platform.system() == 'Darwin':       # macOS
        subprocess.call(('open', doc_filename))
    elif platform.system() == 'Windows':    # Windows
        os.startfile(doc_filename)
    elif platform.system() == 'Linux':      # Linux
        subprocess.call(('xdg-open', doc_filename))


def check_for_files():
    """Check if there are any files available for predicting prices."""
    if os.path.exists(DATA_DIR):
        files = [f for f in os.listdir(DATA_DIR) if os.path.isfile(os.path.join(DATA_DIR, f)) and not f.startswith('.')]
        return files
    return []


def main():
    print("\nWelcome to the Financial Analysis Tool!")

    # Check for available files for prediction
    available_files = check_for_files()

    # Ask the user what they want to do
    print("\nWhat would you like to do?")
    print("1: Portfolio Optimization Analysis")
    if available_files:
        print("2: Predict Future Prices (LSTM)")
    else:
        print("2: Predict Future Prices (LSTM) - No available files")

    while True:
        choice = input("\nEnter 1 for Portfolio Optimization or 2 for Predicting Future Prices: ").strip()
        if choice == '1':
            run_portfolio_optimization()
            break
        elif choice == '2' and available_files:
            run_future_price_prediction()
            break
        elif choice == '2' and not available_files:
            print("There are no files available for future price prediction. Please add some data files first.")
        else:
            print("Invalid input. Please try again.")


def run_portfolio_optimization():
    """Run the portfolio optimization process."""
    print("\nStarting Portfolio Optimization Analysis...\n")

    # Ask user for input
    tickers_input = input("Enter the tickers separated by commas (e.g., AAPL,IEF.BOND,BTC.CRYPTO): ").strip()
    if not tickers_input:
        print("Error: No tickers provided. Exiting program.")
        return
    
    tickers = [ticker.strip().upper() for ticker in tickers_input.split(",")]

    # Print the full asset names based on tickers
    print("\nAsset Names:")
    for ticker in tickers:
        asset_name = get_asset_name(ticker)
        print(f"{ticker}: {asset_name}")

    # Ask for start and end dates
    start_date = input("\nEnter the start date (YYYY-MM-DD): ").strip()
    end_date = input("Enter the end date (YYYY-MM-DD): ").strip()

    # Fetch data
    print("\nFetching data, please wait...")
    data = fetch_data(tickers, start_date, end_date)

    # Ask for a filename to save the data
    filename = input("\nEnter the filename to save the data (e.g., 'portfolio_data.csv'): ").strip()
    save_data(data, filename)
    print(f"Data saved to {filename}.\n")

    # Calculate daily returns
    returns = data.pct_change().dropna()

    # Ask the user for weights
    print("Please enter the weights for each asset in your portfolio (must sum to 1):")
    weights = []
    for ticker in tickers:
        while True:
            try:
                weight = float(input(f"Enter the weight for {ticker} (e.g., 0.2 for 20%): "))
                weights.append(weight)
                break
            except ValueError:
                print("Invalid input. Please enter a valid number.")
    
    # Ensure the weights sum to 1
    if not np.isclose(sum(weights), 1.0):
        print("Error: The sum of the weights must be exactly 1. Please restart the program.")
        return

    # Calculate portfolio metrics
    weights = np.array(weights)
    print("\nCalculating portfolio metrics, please wait...")
    metrics = calculate_portfolio_metrics(weights, returns)
    
    # Display portfolio metrics
    print("\nPortfolio Metrics:")
    print(f"  - Expected Return: {metrics['return'] * 100:.2f}%")
    print(f"  - Volatility: {metrics['volatility'] * 100:.2f}%")
    print(f"  - Sharpe Ratio: {metrics['sharpe_ratio']:.2f}")

    # Optimize portfolio
    print("\nOptimizing portfolio...")
    optimization_result = optimize_portfolio(returns, method='sharpe')
    optimal_weights = optimization_result['weights']
    optimal_metrics = optimization_result['metrics']

    # Display optimized portfolio results
    print("\nOptimized Portfolio Metrics:")
    print(f"  - Optimal Weights:")
    for ticker, weight in zip(tickers, optimal_weights):
        print(f"    {ticker}: {weight:.2%}")
    
    print(f"  - Optimized Expected Return: {optimal_metrics['return']*100:.2f}%")
    print(f"  - Optimized Volatility: {optimal_metrics['volatility']*100:.2f}%")
    print(f"  - Optimized Sharpe Ratio: {optimal_metrics['sharpe_ratio']:.2f}")

    # Save results to a JSON file with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    json_filename = os.path.join(USER_LOG_DIR, f"portfolio_results_{timestamp}.json")
    results = {
        "tickers": tickers,
        "start_date": start_date,
        "end_date": end_date,
        "initial_metrics": metrics,
        "optimal_weights": optimal_weights.tolist(),
        "optimal_metrics": optimal_metrics
    }
    with open(json_filename, 'w') as f:
        json.dump(results, f)
    print(f"\nResults saved to {json_filename}")

    # Ask if the user wants to generate a Word document
    create_doc = input("\nWould you like to generate a Word document with the results? (yes/no): ").strip().lower()
    if create_doc == 'yes':
        doc_filename = os.path.join(USER_LOG_DIR, f"portfolio_results_{timestamp}.docx")
        create_word_doc(results, doc_filename)
        print(f"\nWord document saved to {doc_filename}")
        open_word_doc(doc_filename)


def run_future_price_prediction():
    """Function to load historical data and predict future prices for each ticker."""
    
    # Check if DATA_DIR exists
    if not os.path.exists(DATA_DIR):
        print(f"The directory {DATA_DIR} does not exist.")
    else:
        # List available files in the directory, excluding hidden files
        files = [f for f in os.listdir(DATA_DIR) if os.path.isfile(os.path.join(DATA_DIR, f)) and not f.startswith('.')]
        
        if not files:
            print("No CSV files found in the directory.")
        else:
            # Display available files and prompt the user to select one
            print("Available files:")
            for idx, file in enumerate(files):
                print(f"{idx}: {file}")
            
            # Input validation for selecting a file
            while True:
                try:
                    file_idx = int(input("Select the file number you want to use: "))
                    if 0 <= file_idx < len(files):
                        break
                    else:
                        print(f"Please enter a number between 0 and {len(files) - 1}.")
                except ValueError:
                    print("Invalid input. Please enter a valid number.")
            
            selected_file = files[file_idx]
            
            historical_data, tickers = load_historical_data(selected_file)
            
            #RIGHT HERE
            if historical_data is not None:
                # Fetch and print exchange information for each ticker
                for ticker in tickers:
                    print(f"\nFetching exchange information for ticker: {ticker}")
                    get_exchange_for_ticker(ticker)
                
                # Ask the user how many days to predict
                future_days = int(input("\nEnter the number of days to predict: "))
                
                # Loop through each column (asset) in the data
                for column in historical_data.columns:
                    print(f"\nProcessing asset: {column}")
                    
                    model = build_lstm_model((60, 1))
                    trained_model, scaler, predictions, y_test = train_lstm_model(historical_data[[column]], model)
                    
                    # Plotting the actual vs predicted prices
                    plt.figure(figsize=(12,6))
                    plt.plot(y_test, color='blue', label='Actual Stock Price')
                    plt.plot(predictions, color='red', label='Predicted Stock Price')
                    plt.title(f'Stock Price Prediction for {column}')
                    plt.xlabel('Time')
                    plt.ylabel('Stock Price')
                    plt.legend()
                    plt.show()
                    
                    future_prices = predict_future_prices_lstm(trained_model, historical_data[[column]], scaler, column, future_days=future_days)
                    print(f"Future prices for {column}:")
                    print(future_prices)
                    
                    # Ensure the output directory exists
                    os.makedirs(PREDICTED_DIR, exist_ok=True)
                    
                    # Save the future prices to a CSV file in the specified directory
                    output_file = os.path.join(PREDICTED_DIR, f"{column}_future_prices.csv")
                    future_prices.to_csv(output_file)
                    print(f"Predicted prices saved to {output_file}")

if __name__ == "__main__":
    main()
